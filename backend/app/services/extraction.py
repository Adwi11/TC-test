from __future__ import annotations
from dataclasses import dataclass, field
from typing import Literal
import io
import zipfile
import fitz  # pymupdf
from docx import Document as DocxDocument


Route = Literal["text", "ocr", "vision"]

MIN_CHARS_PER_PAGE = 80
IMAGE_AREA_RATIO_TRIGGER = 0.55
DOCX_IMAGE_EXT = (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp")


@dataclass
class ExtractionStats:
    """Heuristic stats used to pick an extraction route."""

    text: str
    pages: int
    chars_per_page: float
    image_area_ratio: float
    docx_images: list[bytes] = field(default_factory=list)


def _pdf_image_area_ratio(doc: fitz.Document) -> float:
    """Return ratio of image area to page area, averaged across pages."""
    if doc.page_count == 0:
        return 0.0
    ratios: list[float] = []
    for page in doc:
        page_area = max(page.rect.width * page.rect.height, 1.0)
        img_area = 0.0
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                rects = page.get_image_rects(xref)
                for r in rects:
                    img_area += r.width * r.height
            except Exception:
                continue
        ratios.append(min(img_area / page_area, 1.0))
    return sum(ratios) / len(ratios) if ratios else 0.0


def extract_pdf(data: bytes) -> ExtractionStats:
    """Extract text + heuristic stats from a PDF byte string."""
    with fitz.open(stream=data, filetype="pdf") as doc:
        pages = doc.page_count or 1
        text = "\n".join(page.get_text("text") for page in doc)
        chars_per_page = len(text) / pages
        image_ratio = _pdf_image_area_ratio(doc)
    return ExtractionStats(text=text, pages=pages, chars_per_page=chars_per_page, image_area_ratio=image_ratio)


def _docx_media_images(data: bytes) -> list[bytes]:
    """Return raw image bytes from a DOCX's word/media/ directory."""
    images: list[bytes] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/") and name.lower().endswith(DOCX_IMAGE_EXT):
                    try:
                        images.append(zf.read(name))
                    except Exception:
                        continue
    except Exception:
        return []
    return images


def extract_docx(data: bytes) -> ExtractionStats:
    """Extract text + embedded images from a DOCX byte string."""
    doc = DocxDocument(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells if c.text)
    text = "\n".join(parts)
    images = _docx_media_images(data)
    return ExtractionStats(
        text=text,
        pages=1,
        chars_per_page=float(len(text)),
        image_area_ratio=0.0,
        docx_images=images,
    )


def route_decision(stats: ExtractionStats) -> Route:
    """Decide which extraction path to use based on heuristics."""
    if stats.chars_per_page < MIN_CHARS_PER_PAGE:
        return "ocr"
    if stats.image_area_ratio >= IMAGE_AREA_RATIO_TRIGGER:
        return "vision"
    return "text"
